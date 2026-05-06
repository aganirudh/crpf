"""Word document parser — text + table cells with synthetic bboxes.

Word doesn't carry coordinates; we synthesise (page=1, bbox=(0, line_idx*15,
600, line_idx*15+15)) so the downstream provenance contract still holds.
The frontend renders Word docs as text rather than overlaying bboxes.
"""

from __future__ import annotations

import io
import logging

import docx

from pramaan.ingestion.router import PageBlock, hash_text

log = logging.getLogger(__name__)


def parse_docx(data: bytes) -> list[PageBlock]:
    blocks: list[PageBlock] = []
    try:
        doc = docx.Document(io.BytesIO(data))
    except Exception as exc:
        log.warning("docx parse failed: %s", exc)
        return blocks

    line = 0
    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        line += 1
        blocks.append(
            PageBlock(
                page=1,
                text=text,
                bbox=(0.0, float(line * 15), 600.0, float(line * 15 + 15)),
                ocr_conf=1.0,
                extractor="python-docx",
                source_text_sha256=hash_text(text),
            )
        )

    for table in doc.tables:
        for row in table.rows:
            cells_text = [(cell.text or "").strip() for cell in row.cells]
            row_text = " | ".join(t for t in cells_text if t)
            if not row_text:
                continue
            line += 1
            blocks.append(
                PageBlock(
                    page=1,
                    text=row_text,
                    bbox=(0.0, float(line * 15), 600.0, float(line * 15 + 15)),
                    ocr_conf=1.0,
                    extractor="python-docx:table",
                    source_text_sha256=hash_text(row_text),
                    extra={"table_cells": cells_text},
                )
            )

    return blocks
